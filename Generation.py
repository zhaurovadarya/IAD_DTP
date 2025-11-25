import os
import sys
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt

def load_csv_checked(path):
    if not os.path.exists(path):
        print(f"ERROR: Файл не найден: {path}")
        return None
    try:
        df = pd.read_csv(path, low_memory=False)
        print(f"Loaded: {path} ({df.shape[0]} rows, {df.shape[1]} cols)")
        return df
    except Exception as e:
        print(f"ERROR: Не удалось загрузить {path}: {e}")
        return None

def load_data():
    assoc = load_csv_checked("PDD_coincidence.csv")
    bayes = load_csv_checked("Bayes_importance.csv")
    fuzzy = load_csv_checked("Fuzzy_PDD.csv")
    if assoc is None or bayes is None or fuzzy is None:
        print("Один или несколько входных файлов отсутствуют или не загрузились.")
        return None, None, None
    return assoc, bayes, fuzzy

def create_recommendations(assoc: pd.DataFrame, bayes: pd.DataFrame, fuzzy: pd.DataFrame):
    recommendations = []
    if not assoc.empty:
        grouped = assoc.groupby('pdd_id', sort=False)
        for pdd_id, group in grouped:
            pdd_text = group['pdd_text'].iloc[0] if 'pdd_text' in group.columns else ""
            themes = group['themes'].iloc[0] if 'themes' in group.columns else ""
            matched = sorted({f.strip() for s in group['matched_factors'].astype(str) for f in s.split(',') if f.strip()})
            matched_str = ", ".join(matched) if matched else "—"
            recommendations.append({
                "pdd_id": pdd_id,
                "source": "Ассоциативные правила",
                "text": (
                    f"Пункт ПДД {pdd_id}: {pdd_text[:200]}...\n"
                    f"Связанные факторы: {matched_str}. "
                    f"Рекомендация: уточнить формулировку пункта или добавить примеры/требования, "
                    f"учитывая перечисленные факторы. (Темы: {themes})")})

    if not bayes.empty:
        if 'factor' in bayes.columns and 'mean_entropy' in bayes.columns:
            top_bayes = bayes.sort_values("mean_entropy").head(15)
            for _, row in top_bayes.iterrows():
                recommendations.append({
                    "pdd_id": "-",
                    "source": "Байесовская сеть",
                    "text": (
                        f"Фактор «{row['factor']}» обладает низкой средней энтропией ({row['mean_entropy']:.4f}), "
                        f"т.е. информативен для определения тяжести ДТП. Рекомендуется проверить покрытие "
                        f"этого фактора в соответствующих статьях ПДД и, при необходимости, усилить требования.")})
        else:
            print("WARNING: В Bayes_importance.csv отсутствуют ожидаемые столбцы 'factor'/'mean_entropy'.")

    if not fuzzy.empty:
        if 'pdd_id' in fuzzy.columns and 'severity_score_fuzzy' in fuzzy.columns:
            fuzzy_sorted = fuzzy.sort_values("severity_score_fuzzy", ascending=False).head(50)
            for _, row in fuzzy_sorted.iterrows():
                pdd_id = row.get('pdd_id', '-')
                sev = row.get('severity_score_fuzzy', None)
                cat = row.get('severity_category', None)
                related = row.get('related_factors', "")
                recommendations.append({
                    "pdd_id": pdd_id,
                    "source": "Нечеткая логика",
                    "text": (
                        f"Пункт ПДД {pdd_id} получил модельную оценку тяжести {sev} "
                        f"({cat}). Рекомендуется рассмотреть усиление мер контроля или "
                        f"добавление требований по факторам: {related}."
                    )
                })
        else:
            print("WARNING: В Fuzzy_PDD.csv отсутствуют ожидаемые столбцы ('pdd_id','severity_score_fuzzy').")
    return recommendations

def save_docx(recommendations, filename="PDD_Recomd.docx"):
    if not recommendations:
        print("No recommendations to save.")
        return False
    doc = Document()
    title = doc.add_heading("Предложения по внесению изменений в ПДД", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0,0,0)  # Чёрный
        run.font.bold = True
        run.font.size = Pt(16)

    sections = [
        ("Ассоциативные правила", "Ассоциативные правила"),
        ("Байесовские сети", "Байесовская сеть"),
        ("Нечеткая логика", "Нечеткая логика")]

    for sec_title, source_key in sections:
        h = doc.add_heading(sec_title, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.bold = True
            run.font.size = Pt(12)
        items = [r for r in recommendations if r['source'] == source_key]
        if not items:
            doc.add_paragraph("Нет рекомендаций из данного источника.")
            continue

        for idx, rec in enumerate(items, start=1):
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run_num = p.add_run(f"{idx}) ")
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(0,0,0)

            run_head = p.add_run(f"PDD {rec['pdd_id']}: ")
            run_head.font.bold = True
            run_head.font.color.rgb = RGBColor(0,0,0)

            run_text = p.add_run(rec['text'])
            run_text.font.color.rgb = RGBColor(0,0,0)

    # Сохранение
    try:
        doc.save(filename)
        print(f"Документ успешно сохранён: {filename}")
        return True
    except Exception as e:
        print(f"ERROR: Не удалось сохранить документ: {e}")
        return False

def main():
    assoc, bayes, fuzzy = load_data()
    if assoc is None or bayes is None or fuzzy is None:
        sys.exit(1)

    recs = create_recommendations(assoc, bayes, fuzzy)
    print(f"Сформировано рекомендаций: {len(recs)}")
    ok = save_docx(recs)
    if not ok:
        sys.exit(2)

if __name__ == "__main__":
    main()
