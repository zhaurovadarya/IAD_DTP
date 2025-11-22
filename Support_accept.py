import pickle
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor

# Загружаем результаты
with open("analysis_results.pkl", "rb") as f:
    results = pickle.load(f)

# --- Вспомогательная функция для заголовков ---
def add_heading(doc, text, level=1):
    """Добавляет заголовок с черным цветом; Heading 1 — полужирный"""
    p = doc.add_heading(text, level=level)
    if p.runs:
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)  # черный цвет
        if level == 1:
            run.font.bold = True
    return p

# --- Функция для добавления таблиц ---
def add_df_table(doc, df):
    """Надёжно добавляет DataFrame в DOCX, включая категориальные столбцы и NaN"""
    df = df.copy()
    for col in df.columns:
        df[col] = df[col].astype(str)
    df = df.fillna("")

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells

    # Заголовки
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    # Данные
    for i in range(df.shape[0]):
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            value = df.iloc[i, j]
            row_cells[j].text = str(value) if value is not None else ""

# --- Основная функция создания отчета ---
def create_report(results):

    doc = Document()
    add_heading(doc, "Итоговый отчет анализа ДТП", level=1)

    # =======================================================
    # 1. Корреляционный анализ
    # =======================================================
    add_heading(doc, "1. Корреляционный анализ", level=1)
    cor = results.get('correlation', {})

    if cor:
        # --- ГОРОД ---
        add_heading(doc, "1.1. Городские дороги", level=2)
        gor_plot = cor.get("gor_plot")
        gor_info = cor.get("gor_interpretation")
        doc.add_paragraph(
            "Матрица корреляции представлена на графике ниже. "
            "Ключевые коэффициенты:"
        )
        if gor_info:
            doc.add_paragraph(
                f"• Корреляция между плотностью потока и ДТП: r = {gor_info['r_density']:.2f}\n"
                f"• Корреляция между дистанцией между объектами и ДТП: r = {gor_info['r_dist']:.2f}"
            )
            doc.add_paragraph(
                f"Среднее ДТП: {gor_info['mean_acc']:.2f}, std ДТП: {gor_info['std_acc']:.2f}\n"
                f"Δплотности: -7.50 → ΔДТП: {gor_info['delta_acc_density']:+.2f} шт. "
                f"(~{gor_info['percent_density']:.1f}%)\n"
                f"Δдистанции: -2.0 м → ΔДТП: {gor_info['delta_acc_dist']:+.2f} шт. "
                f"(~{gor_info['percent_dist']:.1f}%)"
            )
        if gor_plot:
            doc.add_picture(gor_plot, width=Inches(6))

        # --- ОБЛАСТЬ ---
        add_heading(doc, "1.2. Областные дороги", level=2)
        obl_plot = cor.get("obl_plot")
        obl_info = cor.get("obl_interpretation")
        doc.add_paragraph(
            "Матрица корреляции представлена на графике ниже. "
            "Ключевые коэффициенты:"
        )
        if obl_info:
            doc.add_paragraph(
                f"• Корреляция между трафиком наружу и ДТП: r = {obl_info['r_outward']:.2f}\n"
                f"• Корреляция между трафиком внутрь и ДТП: r = {obl_info['r_return']:.2f}"
            )
            doc.add_paragraph(
                f"Среднее ДТП: {obl_info['mean_acc']:.2f}, std ДТП: {obl_info['std_acc']:.2f}\n"
                f"Δтрафика наружу: +1000 авто → ΔДТП: {obl_info['delta_acc_outward']:+.2f} шт. "
                f"(~{obl_info['percent_outward']:.1f}%)\n"
                f"Δтрафика внутрь: +500 авто → ΔДТП: {obl_info['delta_acc_return']:+.2f} шт. "
                f"(~{obl_info['percent_return']:.1f}%)"
            )
        if obl_plot:
            doc.add_picture(obl_plot, width=Inches(6))

    # =======================================================
    # 2. Логистическая регрессия – город
    # =======================================================
    add_heading(doc, "2. Логистическая регрессия (городские наблюдения)", level=1)
    log_city = results.get("log_city", {}).get("log_city", {})

    add_heading(doc, "2.1. Метрики классификации", level=2)
    class_report = pd.DataFrame([x.split() for x in log_city["classification_report"].split("\n") if x])
    add_df_table(doc, class_report)

    add_heading(doc, "2.2. Матрица ошибок", level=2)
    conf_df = pd.DataFrame(log_city["conf_matrix"], index=[0, 1], columns=[0, 1])
    conf_df.index.name = "Истинный класс"
    conf_df.columns.name = "Предсказанный класс"
    add_df_table(doc, conf_df)

    add_heading(doc, "2.3. ROC-кривая", level=2)
    doc.add_paragraph(f"ROC AUC = {log_city['roc_auc']:.3f}")
    doc.add_picture(log_city["roc_plot"], width=Inches(6))

    add_heading(doc, "2.4. Коэффициенты признаков и интерпретация", level=2)
    coef_df = log_city["coef_df"].copy()
    add_df_table(doc, coef_df)
    doc.add_picture(log_city["coef_plot"], width=Inches(6))
    for i, row in coef_df.iterrows():
        effect = "повышает" if row['Коэффициент'] > 0 else "снижает"
        doc.add_paragraph(
            f"!Признак '{row['Признак']}' {effect} вероятность ДТП примерно в {row['Отношение шансов (e^coef)']:.2f} раз"
        )

    add_heading(doc, "2.5. Распределение предсказанных вероятностей ДТП", level=2)
    doc.add_picture(log_city["prob_dist_plot"], width=Inches(6))

    # =======================================================
    # 3. Логистическая регрессия – область
    # =======================================================
    add_heading(doc, "3. Логистическая регрессия (областные наблюдения)", level=1)
    log_region = results.get("log_region", {}).get("log_region", {})

    add_heading(doc, "3.1. Метрики классификации", level=2)
    class_report = pd.DataFrame([x.split() for x in log_region["classification_report"].split("\n") if x])
    add_df_table(doc, class_report)

    add_heading(doc, "3.2. Матрица ошибок", level=2)
    conf_df = pd.DataFrame(log_region["conf_matrix"], index=[0, 1], columns=[0, 1])
    conf_df.index.name = "Истинный класс"
    conf_df.columns.name = "Предсказанный класс"
    add_df_table(doc, conf_df)

    add_heading(doc, "3.3. ROC-кривая", level=2)
    doc.add_paragraph(f"ROC AUC = {log_region['roc_auc']:.3f}")
    doc.add_picture(log_region["roc_plot"], width=Inches(6))

    add_heading(doc, "3.4. Коэффициенты признаков и интерпретация", level=2)
    coef_df = log_region["coef_df"].copy()
    add_df_table(doc, coef_df.head(10))
    doc.add_picture(log_region["coef_plot"], width=Inches(6))
    for i, row in coef_df.iterrows():
        if abs(row['Коэффициент']) < 0.1:
            continue
        effect = "повышает" if row['Коэффициент'] > 0 else "снижает"
        doc.add_paragraph(
            f"!Признак '{row['Признак']}' {effect} вероятность ДТП примерно в {row['Отношение шансов (e^coef)']:.2f} раз"
        )

    add_heading(doc, "3.5. Распределение предсказанных вероятностей ДТП", level=2)
    doc.add_picture(log_region["prob_dist_plot"], width=Inches(6))

    # =======================================================
    # 4. Модели случайного леса
    # =======================================================
    add_heading(doc, "4. Модели случайного леса", level=1)
    rf_block = results.get("random_forest", {})
    rf_dict = rf_block.get("rf", {})
    geo_df = rf_block.get("geo", pd.DataFrame())
    report_files = rf_block.get("report_files", {})

    # Словарь читаемых названий моделей
    rf_name_map = {
        "rf_type": "Классификация вида ДТП",
        "rf_sev": "Регрессия тяжести ДТП",
        "rf_bin": "Бинарная классификация тяжести ДТП"
    }
    # Top-10 признаков
    model_keys = ["rf_type", "rf_sev", "rf_bin"]
    for key in model_keys:
        df_rf = rf_dict.get(key)
        if df_rf is not None and isinstance(df_rf, pd.DataFrame):
            display_name = rf_name_map.get(key, key)  # читаемое название
            add_heading(doc, f"Random Forest: {display_name}", level=2)
            add_df_table(doc, df_rf.head(10))
            img_file = report_files.get(key)
            if img_file:
                doc.add_picture(img_file, width=Inches(6))

    # Диаграммы географического распределения
    add_heading(doc, "Диаграммы географического распределения", level=2)
    geo_plots = ["boxplot_type", "city_bar", "region_bar"]
    for plot_key in geo_plots:
        img_file = report_files.get(plot_key)
        if img_file:
            doc.add_picture(img_file, width=Inches(6))

    # Географические данные
    if isinstance(geo_df, pd.DataFrame) and not geo_df.empty:
        add_heading(doc, "Географические данные (первые строки)", level=2)
        add_df_table(doc, geo_df.head())
        doc.add_paragraph(
            "Интерактивная карта ДТП сохранена в файле DTP_heatmap.html. "
            "Её можно открыть в браузере для просмотра тепловой карты."
        )

    # =======================================================
    # СОХРАНЕНИЕ
    # =======================================================
    doc.save("Analysis_Report.docx")
    print("✅ Итоговый отчет создан: Analysis_Report.docx")
    return doc

# Запуск
create_report(results)

